
import os
import re
import base64
import tempfile
import requests
from io import BytesIO
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_bookmark(paragraph, bookmark_name):
    """
    Insert a bookmark start and end in the paragraph.
    This allows internal linking (TOC, cross-refs).
    """
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0') # ID needs to be unique in real doc, but 0 often works for simple cases or needs management
    start.set(qn('w:name'), bookmark_name)
    tag.append(start)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    tag.append(end)

def html_to_docx(doc, html_content, theme=0):
    """
    Parses HTML content and adds elements to the python-docx Document.
    Handles Headings, Paragraphs, Lists, Images, and custom 'Note' tables.
    """
    if not html_content: return

    soup = BeautifulSoup(html_content, 'html.parser')

    # Determine base text color for theme
    def get_theme_rgb():
        if theme == 1: return RGBColor(255, 255, 255)
        if theme == 2: return RGBColor(67, 52, 34)
        if isinstance(theme, str) and theme.startswith("#"):
            brightness = get_brightness(theme)
            return RGBColor(255, 255, 255) if brightness < 128 else RGBColor(0, 0, 0)
        return RGBColor(0, 0, 0)
    
    base_rgb = get_theme_rgb()

    def process_node(node, parent_container=None):
        container = parent_container if parent_container else doc

        if isinstance(node, Tag):
            tag = node.name.lower()
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1])
                text = node.get_text(strip=True)
                if text:
                    h = container.add_heading(text, level=level)
                    if theme != 0:
                        for run in h.runs: run.font.color.rgb = base_rgb
            
            elif tag == 'p':
                p = container.add_paragraph()
                process_inline_content(p, node)
                
                # Check for alignment
                style = node.get('style', '')
                if 'text-align: center' in style:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'text-align: right' in style:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'text-align: justify' in style:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            elif tag in ['ul', 'ol']:
                style = 'List Bullet' if tag == 'ul' else 'List Number'
                for li in node.find_all('li', recursive=False):
                    p = container.add_paragraph(style=style)
                    process_inline_content(p, li)

            elif tag == 'img':
                handle_image(container, node)

            elif tag == 'table':
                handle_table(container, node)
            
            elif tag == 'div':
                 # Treat div as container, just recurse
                 for child in node.children:
                     process_node(child, container)
            
            elif tag == 'hr':
                p = container.add_paragraph()
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                if theme == 1:
                    bottom.set(qn('w:color'), 'FFFFFF')
                elif theme == 2:
                    bottom.set(qn('w:color'), 'DCD1BC')
                elif isinstance(theme, str) and theme.startswith("#"):
                    # Use text color for HR in custom theme
                    bottom.set(qn('w:color'), 'FFFFFF' if get_brightness(theme) < 128 else '444444')
                else:
                    bottom.set(qn('w:color'), 'auto')
                pBdr.append(bottom)
                pPr.append(pBdr)

    def process_inline_content(paragraph, html_element):
        for child in html_element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    if theme != 0:
                        run.font.color.rgb = base_rgb
            elif isinstance(child, Tag):
                tag = child.name.lower()
                if tag == 'br':
                    paragraph.add_run().add_break()
                elif tag == 'img':
                    handle_image(doc, child) 
                else:
                    run = paragraph.add_run()
                    run.text = child.get_text() 
                    
                    if tag in ['b', 'strong']:
                        run.bold = True
                    if tag in ['i', 'em']:
                        run.italic = True
                    if tag in ['u', 'ins']:
                        run.underline = True
                    if tag in ['s', 'del', 'strike']:
                        run.font.strike = True
                    
                    style = child.get('style', '')
                    color_match = re.search(r'color:\s*#([0-9a-fA-F]{6})', style)
                    if color_match:
                         hex_color = color_match.group(1)
                         run.font.color.rgb = RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:], 16))
                    elif theme != 0:
                         run.font.color.rgb = base_rgb

    def handle_image(container, img_tag):
        src = img_tag.get('src')
        if not src: return
        try:
            image_stream = None
            if src.startswith('data:image'):
                header, encoded = src.split(',', 1)
                data = base64.b64decode(encoded)
                image_stream = BytesIO(data)
            elif src.startswith('http'):
                resp = requests.get(src, timeout=5)
                if resp.status_code == 200:
                   image_stream = BytesIO(resp.content)
            if image_stream:
                container.add_picture(image_stream, width=Inches(5.0))
        except Exception as e:
            p = container.add_paragraph(f"[Image: {src[:20]}...]")
            if theme != 0:
                for run in p.runs: run.font.color.rgb = base_rgb

    def handle_table(container, table_tag):
        rows = table_tag.find_all('tr')
        if not rows: return
        num_rows = len(rows)
        max_cols = 0
        for r in rows:
            cols = r.find_all(['td', 'th'])
            max_cols = max(max_cols, len(cols))
        if max_cols == 0: return

        docx_table = container.add_table(rows=num_rows, cols=max_cols)
        docx_table.style = 'Table Grid'
        style = table_tag.get('style', '')
        bg_color = None
        if 'background-color' in style:
             match = re.search(r'background-color:\s*#([0-9a-fA-F]{6})', style)
             if match: bg_color = match.group(1).upper()

        for i, row in enumerate(rows):
            cols = row.find_all(['td', 'th'])
            for j, col in enumerate(cols):
                if j >= max_cols: break
                cell = docx_table.cell(i, j)
                cell._element.clear_content()
                for child in col.children:
                    process_node(child, cell)
                
                if bg_color:
                    set_cell_background(cell, bg_color)
                elif theme == 1:
                    set_cell_background(cell, "2D2D2D")
                elif theme == 2:
                    set_cell_background(cell, "EDE6D9")
                elif isinstance(theme, str) and theme.startswith("#"):
                    # Lighten or darken cell based on background
                    if get_brightness(theme) < 128:
                        set_cell_background(cell, "333333")
                    else:
                        set_cell_background(cell, "EEEEEE")

    for child in soup.body.children if soup.body else soup.children:
        process_node(child)


def set_cell_background(cell, hex_color):
    """
    Set background color of a table cell.
    """
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_docx_background(doc, fill_color):
    """
    Sets the page background color for the Word document.
    """
    # 1. Add background element to document
    shd = OxmlElement('w:background')
    shd.set(qn('w:color'), fill_color)
    # Removing themeColor as it overrides the custom hex color
    # shd.set(qn('w:themeColor'), 'background1')
    
    # Clear background
    shd_child = OxmlElement('w:shd')
    shd_child.set(qn('w:val'), 'clear')
    shd_child.set(qn('w:color'), 'auto')
    shd_child.set(qn('w:fill'), fill_color)
    shd.append(shd_child)
    
    doc.element.insert(0, shd)
    
    # 2. Tell Word to display background shapes (crucial for visibility)
    settings = doc.settings.element
    display_bg = settings.find(qn('w:displayBackgroundShape'))
    if display_bg is None:
        display_bg = OxmlElement('w:displayBackgroundShape')
        settings.append(display_bg)

def get_brightness(hex_color):
    """Calculate relative brightness of a hex color (0-255)."""
    hex_color = hex_color.lstrip('#')
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return (r * 299 + g * 587 + b * 114) / 1000

def export_note_to_docx(note, output_path, theme=0):
    doc = Document()
    
    rgb_text = RGBColor(0, 0, 0)
    # Apply background color
    if theme == 1:
        set_docx_background(doc, "1E1E1E")
        rgb_text = RGBColor(255, 255, 255)
    elif theme == 2:
        set_docx_background(doc, "F5F0E8")
        rgb_text = RGBColor(67, 52, 34)
    elif isinstance(theme, str) and theme.startswith("#"):
        set_docx_background(doc, theme.lstrip('#'))
        brightness = get_brightness(theme)
        # Choose text color based on background brightness
        rgb_text = RGBColor(0, 0, 0) if brightness > 128 else RGBColor(255, 255, 255)
    
    if theme != 0:
        style = doc.styles['Normal']
        style.font.color.rgb = rgb_text
    
    # Title
    h = doc.add_heading(note.title, 0)
    if theme != 0:
        for run in h.runs: run.font.color.rgb = rgb_text
    
    # Metadata
    p = doc.add_paragraph()
    run = p.add_run(f"Created: {note.created_at}")
    run.italic = True
    run.font.size = Pt(9)
    if theme == 1:
        run.font.color.rgb = RGBColor(200, 200, 200)
    elif theme == 2:
        run.font.color.rgb = RGBColor(142, 92, 46)
    elif theme != 0:
        # Subtle metadata color for custom theme
        if rgb_text == RGBColor(255, 255, 255):
            run.font.color.rgb = RGBColor(200, 200, 200)
        else:
            run.font.color.rgb = RGBColor(100, 100, 100)
    else:
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph() # Spacer
    
    # Content
    html_to_docx(doc, note.content, theme=theme)
    
    doc.save(output_path)
    return True

def export_folder_to_docx(folder, output_path, progress_callback=None, theme=0):
    doc = Document()
    
    rgb_text = RGBColor(0, 0, 0)
    # Apply background color
    if theme == 1:
        set_docx_background(doc, "1E1E1E")
        rgb_text = RGBColor(255, 255, 255)
    elif theme == 2:
        set_docx_background(doc, "F5F0E8")
        rgb_text = RGBColor(67, 52, 34)
    elif isinstance(theme, str) and theme.startswith("#"):
        set_docx_background(doc, theme.lstrip('#'))
        brightness = get_brightness(theme)
        # Choose text color based on background brightness
        rgb_text = RGBColor(0, 0, 0) if brightness > 128 else RGBColor(255, 255, 255)
    
    if theme != 0:
        style = doc.styles['Normal']
        style.font.color.rgb = rgb_text
    
    # Folder Title Page
    h = doc.add_heading(folder.name, 0)
    if theme != 0:
        for run in h.runs: run.font.color.rgb = rgb_text
        
    doc.add_paragraph("Folder Export").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Iterate Notes
    notes = sorted(folder.notes, key=lambda n: n.created_at) # Sort by creation or custom order
    total = len(notes)
    
    for i, note in enumerate(notes):
        if progress_callback: progress_callback(i, total)
        
        h = doc.add_heading(note.title, 1)
        if theme != 0:
            for run in h.runs: run.font.color.rgb = rgb_text
            
        html_to_docx(doc, note.content, theme=theme)
        doc.add_page_break()
        
    doc.save(output_path)
    return True
